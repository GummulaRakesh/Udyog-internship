# import pandas as pd
# from docx import Document
# from docx.shared import Pt
# import comtypes.client

# # Part 1: Excel Data Reading
# def read_excel(file_path, sheet_name):
#     # Load the Excel file
#     df = pd.read_excel(file_path, sheet_name=sheet_name)

#     # Handle missing values (fill NaNs with a default value or raise an alert)
#     df.fillna('Missing', inplace=True)  # You can also use dropna() if needed

#     # Perform basic data validation
#     # Check for duplicate rows
#     if df.duplicated().any():
#         print("Warning: There are duplicate rows in the dataset.")

#     # Check for empty columns
#     if df.isnull().values.any():
#         print("Warning: There are missing values in the dataset.")

#     # Data type validation (checking types of specific columns)
#     # if not pd.api.types.is_numeric_dtype(df['Column1']):  # Replace 'Column1' with actual column name
#     #     print("Error: Column1 contains non-numeric data.")

#     return df
#     # print(df.head(10))

# # Part 2: Word Document Generation
# # def generate_word_report(excel_data, word_template, output_docx, output_pdf):
# #     # Load the Word document template
# #     doc = Document(word_template)

# #     # Replace placeholders in the document with data from Excel
# #     for paragraph in doc.paragraphs:
# #         if 'PLACEHOLDER1' in paragraph.text:  # Replace 'PLACEHOLDER1' with the actual placeholder
# #             paragraph.text = paragraph.text.replace('PLACEHOLDER1', str(excel_data['Column1'][0]))
# #         if 'PLACEHOLDER2' in paragraph.text:
# #             paragraph.text = paragraph.text.replace('PLACEHOLDER2', str(excel_data['Column2'][0]))

# #     # Save the updated document as DOCX
# #     doc.save(output_docx)

# def generate_word_report(excel_data, word_template, output_docx, output_pdf):
#     # Load the Word document template
#     doc = Document(word_template)

#     if not doc.paragraphs:
#         for _ in range(len(excel_data)):  # Add as many paragraphs as rows in the Excel sheet
#             doc.add_paragraph('')

#     # Find specific paragraphs or tables to update based on their index or content.
#     # For example, replace content in the first and second paragraphs with Excel data:
#     paragraphs = doc.paragraphs

#     if len(paragraphs) < len(excel_data):
#         for _ in range(len(excel_data) - len(paragraphs)):
#             doc.add_paragraph('')
    
#     # Modify specific paragraphs with data from Excel
#     # if len(paragraphs) > 0:
#     #     paragraphs[0].text = f"This is new content for the first paragraph: {excel_data['Name'][0]}"
#     # if len(paragraphs) > 1:
#     #     paragraphs[1].text = f"This is new content for the second paragraph: {excel_data['Age'][0]}"
#     # if len(paragraphs) > 2:
#     #     paragraphs[2].text = f"This is new content for the Third paragraph: {excel_data['Location'][0]}"

#     for i in range(0,len(excel_data)):
#             paragraphs[i].text = f"{excel_data['Name'][i]} {excel_data['Age'][i]} {excel_data['Location'][i]}"

#     # if len(paragraphs)>0:
#     #     paragraphs[0].text = f"{excel_data['Name'][0]} {excel_data['Age'][0]} {excel_data['Location'][0]}"

#     # Optionally, modify specific table cells if needed
#     if doc.tables:
#         # Assuming we have tables, and we want to modify the first cell in the first row of the first table
#         table = doc.tables[0]
#         table.cell(0, 0).text = str(excel_data['Column1'][0])

#     # Save the updated document as DOCX
#     doc.save(output_docx)


#     # Export as PDF
#     wdFormatPDF = 17
#     word = comtypes.client.CreateObject('Word.Application')
#     doc = word.Documents.Open(output_docx)
#     doc.SaveAs(output_pdf, FileFormat=wdFormatPDF)
#     doc.Close()
#     word.Quit()

# # Main function to run both parts
# if __name__ == "__main__":
#     # Define file paths
#     excel_file = 'C:\\Users\\admin\\Desktop\\Book1.xlsx'
#     sheet_name = 'Sheet1'
#     word_template = 'C:\\Users\\admin\\Desktop\\Document2.docx'
#     output_docx = 'C:\\Users\\admin\\Desktop\\updated_document.docx'
#     output_pdf = 'C:\\Users\\admin\\Desktop\\output.pdf'

#     # Step 1: Read data from Excel
#     excel_data = read_excel(excel_file, sheet_name)

#     # Step 2: Generate Word report and save as PDF
#     generate_word_report(excel_data, word_template, output_docx, output_pdf)

# import pandas as pd
# from docx import Document
# from docx.shared import Pt
# import comtypes.client

# # Part 1: Excel Data Reading
# def read_excel(file_path, sheet_name):
#     print(f"Reading Excel file: {file_path}")
#     df = pd.read_excel(file_path, sheet_name=sheet_name)
#     print("Excel file read successfully.")

#     # Explicitly cast numeric columns to 'object' (string-compatible) type before filling NaNs
#     for col in df.select_dtypes(include=['float64']).columns:
#         df[col] = df[col].astype('object')
#     print("Converted float columns to object type.")

#     # Handle missing values: fill NaNs with a default value
#     df.fillna('Missing', inplace=True)
#     print("Filled missing values.")

#     # Check for duplicate rows
#     if df.duplicated().any():
#         print("Warning: There are duplicate rows in the dataset.")

#     # Check for empty columns
#     if df.isnull().values.any():
#         print("Warning: There are missing values in the dataset.")

#     return df

# # Part 2: Word Document Generation
# def generate_word_report(excel_data, word_template, output_docx, output_pdf):
#     # Load the Word document template
#     doc = Document(word_template)

#     # Check if the document contains tables
#     if doc.tables:
#         table = doc.tables[0]  # Use the first table in the document

#         # Ensure the table has enough rows for the Excel data
#         while len(table.rows) < len(excel_data):
#             table.add_row()

#         # Iterate over the Excel data and populate the table
#         for i, row in excel_data.iterrows():
#             table.cell(i, 0).text = str(row['Name'])
#             table.cell(i, 1).text = str(row['Age'])
#             table.cell(i, 2).text = str(row['Location'])
#     else:
#         print("No tables found in the document.")

#     # Save the updated document as DOCX
#     doc.save(output_docx)

#     # Convert DOCX to PDF
#     wdFormatPDF = 17
#     word = comtypes.client.CreateObject('Word.Application')
#     doc = word.Documents.Open(output_docx)
#     doc.SaveAs(output_pdf, FileFormat=wdFormatPDF)
#     doc.Close()
#     word.Quit()


#     print(f"PDF document saved: {output_pdf}")

# # Main function to run both parts
# if __name__ == "__main__":
#     excel_file = 'C:\\Users\\admin\\Desktop\\Book1.xlsx'
#     sheet_name = 'Sheet1'
#     word_template = 'C:\\Users\\admin\\Desktop\\Document3.docx'
#     output_docx = 'C:\\Users\\admin\\Desktop\\updated_document.docx'
#     output_pdf = 'C:\\Users\\admin\\Desktop\\output.pdf'

#     print("Starting Excel to Word to PDF script.")

#     # Step 1: Read data from Excel
#     excel_data = read_excel(excel_file, sheet_name)
#     print("Excel data loaded.")

#     # Step 2: Generate Word report and save as PDF
#     generate_word_report(excel_data, word_template, output_docx, output_pdf)
#     print("Process completed successfully.")


import pandas as pd
from docx import Document
from docx.shared import Pt
import comtypes.client


# Part 1: Excel Data Reading
def read_excel(file_path, sheet_name):
    print(f"Reading Excel file: {file_path}")
    df = pd.read_excel(file_path, sheet_name=sheet_name)

    # Handle missing values: fill NaNs with a default value
    df.fillna('Missing', inplace=True)

    # Check for duplicate rows
    if df.duplicated().any():
        print("Warning: There are duplicate rows in the dataset.")

    # Check for empty columns
    if df.isnull().values.any():
        print("Warning: There are missing values in the dataset.")

    return df

# Part 2: Word Document Generation
def generate_word_report(excel_data, output_docx, output_pdf):
    # Create a new Word document
    doc = Document()

    # Set a title for the document
    doc.add_heading('Excel Data Report', level=1)

    # Iterate through Excel data and add content to the Word document
    for i, row in excel_data.iterrows():
        # Add a paragraph for each row of data
        doc.add_paragraph(f"Name: {row['Name']}, Age: {row['Age']}, Location: {row['Location']}, Gender: {row['Gender']}")

    # Save the newly created Word document
    doc.save(output_docx)
    # print(f"Word document saved as: {output_docx}")

    # Convert the Word document to PDF
    convert_to_pdf(output_docx, output_pdf)

# Part 3: Word to PDF Conversion
def convert_to_pdf(input_docx, output_pdf):
    wdFormatPDF = 17
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(input_docx)
    # doc=input_docx
    doc.SaveAs(output_pdf, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()
    print(f"PDF document saved as: {output_pdf}")



# Main function to run both parts
if __name__ == "__main__":
    excel_file = 'C:\\Users\\admin\\Desktop\\Book1.xlsx'
    sheet_name = 'Sheet1'
    # word_template = 'C:\\Users\\admin\\Desktop\\Document3.docx'
    output_docx = 'C:\\Users\\admin\\Desktop\\updated_document.docx'
    output_pdf = 'C:\\Users\\admin\\Desktop\\output.pdf'

    print("Starting Excel to Word to PDF script.")

    # Step 1: Read data from Excel
    excel_data = read_excel(excel_file, sheet_name)
    print("Excel data loaded.")

    # Step 2: Generate Word report and save as PDF
    generate_word_report(excel_data, output_docx, output_pdf)
    print("Process completed successfully.")